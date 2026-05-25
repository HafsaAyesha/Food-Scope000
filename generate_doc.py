from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── PAGE MARGINS ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── HELPERS ───────────────────────────────────────────────────────────────────
DARK  = RGBColor(0x1a, 0x1a, 0x2e)
GOLD  = RGBColor(0xe2, 0xb0, 0x4a)
GREY  = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        val = kwargs.get(edge, {})
        if val:
            tag = OxmlElement(f'w:{edge}')
            for k,v in val.items():
                tag.set(qn(f'w:{k}'), v)
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def heading(text, level=1, color=DARK, size=None, bold=True, space_before=12, space_after=6):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    else:
        sizes = {1: 20, 2: 16, 3: 13, 4: 12}
        run.font.size = Pt(sizes.get(level, 12))
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    return p

def body(text, size=11, color=GREY, space_after=6, italic=False, bold=False):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size       = Pt(size)
    run.font.color.rgb  = color
    run.font.italic     = italic
    run.font.bold       = bold
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(items, indent=True):
    for item in items:
        p   = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size      = Pt(10.5)
        run.font.color.rgb = GREY
        p.paragraph_format.space_after = Pt(3)
        if indent:
            p.paragraph_format.left_indent = Cm(0.5)

def section_label(text):
    p   = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size      = Pt(8)
    run.font.bold      = True
    run.font.color.rgb = GOLD
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(2)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run('─' * 80)
    run.font.size      = Pt(6)
    run.font.color.rgb = RGBColor(0xdd, 0xdd, 0xdd)

# ── COVER PAGE ────────────────────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(48)

run = cover.add_run('FOOD')
run.font.size       = Pt(48)
run.font.bold       = True
run.font.color.rgb  = DARK

run2 = cover.add_run('SCOPE')
run2.font.size      = Pt(48)
run2.font.bold      = True
run2.font.color.rgb = GOLD

sub = doc.add_paragraph('Discover, Review & Explore Restaurants Near You')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.runs[0]
r.font.size      = Pt(13)
r.font.color.rgb = GREY
r.font.italic    = True
sub.paragraph_format.space_after = Pt(36)

# Cover table
tbl = doc.add_table(rows=2, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style     = 'Table Grid'
cells_data = [
    ('Course',          'Advanced Web Technologies'),
    ('Assessment',      'Lab Terminal — MERN Stack Application (CLO5)'),
    ('Technology Stack','MongoDB · Express · React · Node.js'),
    ('Maximum Marks',   '50'),
]
flat = tbl.rows[0].cells + tbl.rows[1].cells
for i, cell in enumerate(flat):
    label, value = cells_data[i]
    set_cell_bg(cell, '1a1a2e')
    p  = cell.paragraphs[0]
    lp = p.add_run(label + '\n')
    lp.font.size      = Pt(8)
    lp.font.color.rgb = GOLD
    lp.font.bold      = True
    vp = p.add_run(value)
    vp.font.size      = Pt(11)
    vp.font.color.rgb = WHITE
    vp.font.bold      = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    cell.width = Cm(8)

doc.add_paragraph()
cfoot = doc.add_paragraph('Project Document  ·  FoodScope Web Application')
cfoot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cfoot.runs[0]
r.font.size = Pt(9); r.font.color.rgb = GREY; r.font.italic = True
cfoot.paragraph_format.space_before = Pt(32)

doc.add_page_break()

# ── TABLE OF CONTENTS ─────────────────────────────────────────────────────────
heading('Table of Contents', level=2, space_before=0)
toc_items = [
    ('1', 'Problem Statement'),
    ('2', 'Proposed Solution'),
    ('3', 'Scope Statement'),
    ('4', 'System Actors'),
    ('5', 'Functional Requirements'),
    ('6', 'Technology Stack'),
    ('7', 'REST API Design'),
    ('8', 'Data Design — Mongoose Schemas'),
    ('9', 'UI Screenshots'),
]
for num, title in toc_items:
    p   = doc.add_paragraph()
    r1  = p.add_run(f'{num}.  ')
    r1.font.bold       = True
    r1.font.color.rgb  = GOLD
    r1.font.size       = Pt(11)
    r2  = p.add_run(title)
    r2.font.size       = Pt(11)
    r2.font.color.rgb  = DARK
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════════════════════
section_label('Section 01')
heading('1. Problem Statement', level=1)

body(
    "Pakistan's food and dining industry is expansive and rapidly growing, yet restaurant discovery "
    "remains fragmented and inefficient. Diners rely on informal word-of-mouth, outdated social media "
    "posts, or generic global platforms that have limited local coverage and no real community engagement "
    "for Pakistani cities such as Karachi, Lahore, and Islamabad."
)
body(
    "Restaurant owners — particularly independent and small-scale operators — have no dedicated, structured "
    "platform to list their establishments, manage their profiles, and receive credible feedback from a "
    "genuine user base. Existing solutions either require expensive subscriptions, are cluttered with "
    "advertisements, or simply do not reflect the local dining landscape accurately."
)

heading('From the Customer\'s Perspective', level=3)
bullet([
    'No centralised platform to browse, filter and compare restaurants by cuisine, price range or rating.',
    'Inability to read and contribute authentic, community-driven reviews.',
    'No reliable geospatial restaurant discovery ("find restaurants near me").',
    'Difficulty identifying dietary options (Halal, Vegetarian, Gluten-Free) when searching.',
])

heading('From the Restaurant Owner\'s Perspective', level=3)
bullet([
    'No self-service portal to create and maintain a digital presence for their restaurant.',
    'No structured way to receive and respond to genuine customer feedback.',
    'Limited visibility to reach new customers beyond their existing social media following.',
])

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 2. PROPOSED SOLUTION
# ══════════════════════════════════════════════════════════════════════════════
section_label('Section 02')
heading('2. Proposed Solution', level=1)

body(
    "FoodScope is a full-stack MERN (MongoDB, Express, React, Node.js) web application that serves as a "
    "community-driven restaurant discovery and review platform tailored for Pakistani cities. It provides "
    "a Single Page Application (SPA) front-end with high interactivity and a robust RESTful back-end API "
    "with clearly defined business logic, JWT authentication, and Role-Based Access Control (RBAC)."
)

heading('Key Pillars of the Solution', level=3)

features = [
    ('Restaurant Discovery',
     'Browse restaurants with multi-criteria filtering (cuisine, price, rating, tags). Geospatial "nearby" search using MongoDB 2dsphere indexes.'),
    ('Community Reviews',
     'Authenticated users post reviews with star ratings and written feedback. A helpful/not-helpful voting system surfaces quality content.'),
    ('Restaurant Owner Portal',
     'Restaurant owners can register and list their establishments with full details — location, cuisine, price range, and menu items.'),
    ('Role-Based Access Control',
     'Three distinct roles (Admin, Restaurant Owner, Customer) with clearly scoped permissions enforced on the API via middleware.'),
    ('Secure Authentication',
     'JWT access + refresh tokens, argon2id password hashing, per-route rate limiting, and account lockout after repeated failed logins.'),
    ('Admin Dashboard',
     'Administrators approve restaurant listings, moderate reviews, manage users, and view platform-wide analytics.'),
]
for title, desc in features:
    p   = doc.add_paragraph()
    r1  = p.add_run(f'  {title}: ')
    r1.font.bold      = True
    r1.font.size      = Pt(11)
    r1.font.color.rgb = DARK
    r2  = p.add_run(desc)
    r2.font.size      = Pt(10.5)
    r2.font.color.rgb = GREY
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Cm(0.5)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 3. SCOPE STATEMENT
# ══════════════════════════════════════════════════════════════════════════════
section_label('Section 03')
heading('3. Scope Statement', level=1)

heading('Business Name & Purpose', level=3)
body(
    "FoodScope — the name reflects the platform's purpose: a wide-angle lens for exploring the food "
    "landscape. FoodScope's online presence exists to connect food lovers with great restaurants and give "
    "restaurant owners a credible, structured digital presence within their local community."
)

heading('Target Market', level=3)
body(
    "Diners and food enthusiasts in Pakistani cities (initial focus: Karachi, Lahore, Islamabad) and "
    "independent restaurant owners seeking online visibility without the cost of large enterprise platforms."
)

heading('Services Offered', level=3)
bullet([
    'Restaurant browsing with multi-criteria filtering (cuisine, price, rating, tags)',
    'Geospatial "nearby restaurants" discovery using GPS coordinates',
    'User account registration, login, and profile management',
    'Review and rating submission with helpful/not-helpful voting',
    'Comment threads on reviews for community discussion',
    'Restaurant bookmarking / favourites for logged-in users',
    'Restaurant owner portal: create and manage listings and dish menus',
    'Tag-based classification (cuisine type, dietary requirements, features)',
    'Admin control panel for content moderation and platform management',
    'In-app notification system for relevant user activity',
])

heading('Out of Scope (Current Version)', level=3)
p   = doc.add_paragraph()
r   = p.add_run(
    'Note: The following features are planned for future iterations and are not part of the current '
    'submission: online reservation/booking system, payment integration, real-time chat between owners '
    'and customers, and a dedicated mobile application.'
)
r.font.size      = Pt(10)
r.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
r.font.italic    = True
p.paragraph_format.left_indent  = Cm(0.5)
p.paragraph_format.space_after  = Pt(6)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 4. SYSTEM ACTORS
# ══════════════════════════════════════════════════════════════════════════════
section_label('Section 04')
heading('4. System Actors', level=1)
body(
    'FoodScope has three distinct user roles, each with specific permissions enforced through '
    'Role-Based Access Control (RBAC) middleware on the API layer.'
)

actors = [
    ('Customer (User)',
     'Registered members of the public who use FoodScope to find and review restaurants.',
     ['Browse and search restaurant listings',
      'Read reviews, ratings and comments',
      'Post, edit and delete their own reviews',
      'Comment on reviews (threaded replies)',
      'Vote reviews helpful or not helpful',
      'Bookmark favourite restaurants',
      'Manage their own profile and avatar']),
    ('Restaurant Owner  [role: reviewer]',
     'Verified businesses who list and manage their restaurant(s) on the platform.',
     ['All Customer permissions',
      'Create new restaurant listings',
      'Edit and update own restaurant details',
      'Add, edit and remove menu dishes',
      'Manage restaurant tags and photos',
      'View restaurant review statistics']),
    ('Administrator',
     'Platform superusers with full control over all content and user accounts.',
     ['All Restaurant Owner permissions',
      'Approve, reject or suspend restaurant listings',
      'Suspend or delete any user account',
      'Moderate any review (hide / archive)',
      'Manage the tag taxonomy system-wide',
      'View platform-wide analytics dashboard']),
]
for name, desc, perms in actors:
    heading(name, level=3, color=DARK, size=12)
    body(desc, size=10.5)
    bullet(perms)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 5. FUNCTIONAL REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════
section_label('Section 05')
heading('5. Functional Requirements', level=1)

fr_groups = [
    ('FR-1: User Authentication & Accounts', [
        'Users can register with name, email, and password',
        'Email verification is required before account activation',
        'Users can log in and receive JWT access and refresh tokens',
        'Forgot password and reset password flows via secure email token',
        'Users can update their profile information and avatar',
        'Account lockout after repeated failed login attempts (brute-force protection)',
    ]),
    ('FR-2: Restaurant Listings', [
        'Restaurant owners can create listings with name, description, cuisine type, price range, address, and geolocation',
        'New listings are submitted as pending and require admin approval before going live',
        'Restaurants can be filtered by cuisine, price range, minimum rating, and tags',
        'Full-text search across restaurant name and description',
        'Geospatial "nearby" search returns restaurants within a configurable kilometre radius',
    ]),
    ('FR-3: Reviews & Ratings', [
        'Authenticated users can submit one review per restaurant (1–5 stars with written body)',
        'Users can edit or delete their own reviews',
        'Other users can vote a review as helpful or not helpful (one vote per user)',
        'Admins can hide or archive inappropriate reviews',
        "Restaurant average rating is automatically recalculated from all active reviews",
    ]),
    ('FR-4: Comments', [
        'Authenticated users can comment on any review',
        'Threaded replies supported up to depth 2',
        'Users can delete their own comments',
    ]),
    ('FR-5: Tags & Classification', [
        'Tags are categorised as cuisine type, dietary requirement, or restaurant feature',
        'Restaurant owners can propose new tags; admins approve or reject them',
        'Tags are used as filters on the restaurant listing page',
    ]),
    ('FR-6: Bookmarks & Notifications', [
        'Logged-in users can bookmark restaurants and view their saved list',
        'In-app notifications for relevant user activity (review responses, etc.)',
    ]),
    ('FR-7: Admin Dashboard', [
        'View and manage all registered users with search and filtering',
        'Approve, reject, or suspend restaurant listings',
        'Moderate reviews (hide/archive with reason)',
        'View platform-wide analytics (user counts, restaurant counts, review counts)',
    ]),
]
for title, items in fr_groups:
    heading(title, level=3, size=12)
    bullet(items)

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 6. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════
section_label('Section 06')
heading('6. Technology Stack', level=1)

tbl = doc.add_table(rows=1, cols=3)
tbl.style     = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for i, txt in enumerate(['Layer', 'Technology', 'Purpose']):
    set_cell_bg(hdr[i], '1a1a2e')
    p   = hdr[i].paragraphs[0]
    r   = p.add_run(txt)
    r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE

stack_rows = [
    ('Database',      'MongoDB Atlas',         'NoSQL document store with 2dsphere geospatial indexes'),
    ('ODM',           'Mongoose',              'Schema definition, validation, and query building'),
    ('Backend',       'Node.js + Express v5',  'RESTful API server with middleware pipeline'),
    ('Frontend',      'React 18',              'Component-based Single Page Application (SPA)'),
    ('Routing',       'React Router v6',       'Client-side navigation with protected routes'),
    ('Build Tool',    'Vite',                  'Frontend bundler with Hot Module Replacement'),
    ('Auth',          'JWT (access + refresh)','Stateless authentication with token rotation'),
    ('Password',      'argon2id',              'Industry-standard password hashing algorithm'),
    ('Validation',    'Joi',                   'Request schema validation on all API routes'),
    ('HTTP Client',   'Axios',                 'Fetch API wrapper with request/response interceptors'),
    ('Security',      'Helmet + CORS',         'HTTP security headers and cross-origin policy'),
    ('Rate Limiting', 'rate-limiter-flexible', 'Per-route rate limiting with in-memory fallback'),
    ('Logging',       'Pino',                  'Structured JSON request and error logging'),
    ('Email',         'Nodemailer',            'Transactional email (password reset, verification)'),
]
for layer, tech, purpose in stack_rows:
    row = tbl.add_row().cells
    for i, txt in enumerate([layer, tech, purpose]):
        p   = row[i].paragraphs[0]
        r   = p.add_run(txt)
        r.font.size = Pt(10)
        if i == 1:
            r.font.bold = True; r.font.color.rgb = DARK
        else:
            r.font.color.rgb = GREY
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

divider()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. REST API DESIGN
# ══════════════════════════════════════════════════════════════════════════════
section_label('Section 07')
heading('7. REST API Design', level=1)
body(
    'All API endpoints are versioned under the base path /api/v1. '
    'Authentication is enforced via Bearer token in the Authorization header. '
    'AUTH denotes endpoints requiring a valid JWT access token.'
)

api_groups = [
    ('Authentication  —  /api/v1/auth', [
        ('POST',   '/register',         'Register a new user account',                              'Public'),
        ('POST',   '/login',            'Authenticate user, return access + refresh tokens',        'Public'),
        ('POST',   '/logout',           'Revoke refresh token and clear session',                   'AUTH'),
        ('POST',   '/refresh',          'Issue a new access token using refresh token',             'Public'),
        ('GET',    '/me',               "Get authenticated user's profile",                         'AUTH'),
        ('PUT',    '/me',               "Update authenticated user's profile",                      'AUTH'),
        ('POST',   '/forgot-password',  'Send password reset link to email',                        'Public'),
        ('POST',   '/reset-password',   'Reset password using valid token',                         'Public'),
    ]),
    ('Restaurants  —  /api/v1/restaurants', [
        ('GET',    '/',                 'List all approved restaurants (with filters & pagination)', 'Public'),
        ('GET',    '/:id',              'Get a single restaurant by ID',                            'Public'),
        ('POST',   '/',                 'Create a new restaurant listing',                          'OWNER'),
        ('PUT',    '/:id',              'Update restaurant details',                                'OWNER'),
        ('DELETE', '/:id',              'Soft-delete a restaurant',                                 'ADMIN'),
        ('POST',   '/:id/bookmark',     'Bookmark / un-bookmark a restaurant',                     'AUTH'),
        ('GET',    '/me/bookmarks',     "Get current user's bookmarked restaurants",               'AUTH'),
    ]),
    ('Reviews  —  /api/v1/reviews', [
        ('GET',    '/',                 'List reviews (filterable by restaurant)',                   'Public'),
        ('POST',   '/',                 'Submit a new review for a restaurant',                     'AUTH'),
        ('PUT',    '/:id',              'Edit own review',                                          'AUTH'),
        ('DELETE', '/:id',              'Delete own review',                                        'AUTH'),
        ('POST',   '/:id/vote',         'Vote a review helpful or not helpful',                    'AUTH'),
        ('GET',    '/:id/comments',     'Get comments on a review',                                'Public'),
        ('POST',   '/:id/comments',     'Add a comment to a review',                               'AUTH'),
    ]),
    ('Search & Geo  —  /api/v1/search & /api/v1/geo', [
        ('GET',    '/search',           'Full-text search across restaurants',                      'Public'),
        ('GET',    '/geo/nearby',       'Find restaurants within a radius of coordinates',          'Public'),
        ('GET',    '/geo/reverse',      'Reverse geocode coordinates to address',                   'Public'),
        ('GET',    '/geo/location/ip',  'Approximate user location from IP address',                'Public'),
    ]),
    ('Tags & Admin  —  /api/v1/tags & /api/v1/admin', [
        ('GET',    '/tags',                          'List all approved tags (filterable by type)',  'Public'),
        ('POST',   '/tags',                          'Create a new tag (pending approval)',          'OWNER'),
        ('GET',    '/admin/users',                   'List all users with filters',                 'ADMIN'),
        ('PATCH',  '/admin/restaurants/:id/status',  'Approve / reject / suspend a restaurant',    'ADMIN'),
        ('PATCH',  '/admin/reviews/:id/moderate',    'Hide or archive a review',                   'ADMIN'),
        ('GET',    '/admin/analytics',               'Platform statistics and analytics',           'ADMIN'),
    ]),
]

METHOD_COLORS = {
    'GET':    '1a7a3c', 'POST':   '1e40af',
    'PUT':    '92400e', 'PATCH':  '5b21b6', 'DELETE': '991b1b',
}
METHOD_BG = {
    'GET':    'dcf5e7', 'POST':   'dbeafe',
    'PUT':    'fef3c7', 'PATCH':  'ede9fe', 'DELETE': 'fee2e2',
}

for group_title, endpoints in api_groups:
    heading(group_title, level=3, size=12, space_before=16)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style     = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, txt in enumerate(['Method', 'Endpoint', 'Description', 'Access']):
        set_cell_bg(hdr[i], '1a1a2e')
        p   = hdr[i].paragraphs[0]
        r   = p.add_run(txt)
        r.font.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = WHITE
    for method, endpoint, desc, access in endpoints:
        row = tbl.add_row().cells
        # Method cell
        set_cell_bg(row[0], METHOD_BG.get(method, 'f0f0f0'))
        pm = row[0].paragraphs[0]
        rm = pm.add_run(method)
        rm.font.bold = True; rm.font.size = Pt(9)
        rm.font.color.rgb = RGBColor.from_string(METHOD_COLORS.get(method, '333333'))
        # Endpoint
        pe = row[1].paragraphs[0]
        re = pe.add_run(endpoint)
        re.font.size = Pt(9); re.font.color.rgb = DARK; re.font.bold = True
        # Description
        pd = row[2].paragraphs[0]
        rd = pd.add_run(desc)
        rd.font.size = Pt(9.5); rd.font.color.rgb = GREY
        # Access
        pa = row[3].paragraphs[0]
        ra = pa.add_run(access)
        ra.font.size = Pt(9); ra.font.bold = True
        if access == 'Public':
            ra.font.color.rgb = RGBColor(0x16, 0x7a, 0x3c)
        elif access == 'ADMIN':
            ra.font.color.rgb = RGBColor(0x99, 0x1b, 0x1b)
        else:
            ra.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
        for cell in row:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after  = Pt(3)
    doc.add_paragraph()

divider()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. DATA DESIGN
# ══════════════════════════════════════════════════════════════════════════════
section_label('Section 08')
heading('8. Data Design — Mongoose Schemas', level=1)
body('All data is stored in MongoDB. The following collections are defined using Mongoose schemas with validation, appropriate indexes, and document references.')

schemas = [
    ('User Collection', [
        ('name',                 'String',    'Required'),
        ('email',                'String',    'Unique index'),
        ('password',             'String',    'argon2id hash'),
        ('role',                 'String',    'enum: user | reviewer | admin'),
        ('isVerified',           'Boolean',   ''),
        ('isSuspended',          'Boolean',   ''),
        ('failedLoginAttempts',  'Number',    'Brute-force protection'),
        ('lockUntil',            'Date',      'Account lockout expiry'),
        ('avatar_url',           'String',    ''),
        ('review_count',         'Number',    'Denormalised counter'),
    ]),
    ('Restaurant Collection', [
        ('name',             'String',    'Required'),
        ('description',      'String',    'Full-text indexed'),
        ('cuisine_type',     'String',    ''),
        ('price_range',      'String',    'enum: $ | $$ | $$$'),
        ('avg_rating',       'Number',    '0–5, auto-updated'),
        ('address',          'String',    ''),
        ('location',         'GeoJSON Point', '2dsphere index'),
        ('tags',             '[String]',  ''),
        ('owner_id',         'ObjectId',  'Ref: User'),
        ('status',           'String',    'enum: pending | approved | rejected'),
        ('thumbnail',        'String',    'Image URL'),
    ]),
    ('Review Collection', [
        ('user_id',              'ObjectId',  'Ref: User'),
        ('restaurant_id',        'ObjectId',  'Ref: Restaurant'),
        ('rating',               'Number',    'Required, 1–5'),
        ('body',                 'String',    'Max 3000 characters'),
        ('helpful_count',        'Number',    ''),
        ('not_helpful_count',    'Number',    ''),
        ('status',               'String',    'enum: active | hidden | archived'),
        ('photos',               '[String]',  'Image URLs'),
    ]),
    ('Comment Collection', [
        ('review_id',            'ObjectId',  'Ref: Review'),
        ('user_id',              'ObjectId',  'Ref: User'),
        ('parent_comment_id',    'ObjectId',  'Ref: Comment (nullable)'),
        ('depth',                'Number',    'Max 2 (threaded replies)'),
        ('body',                 'String',    'Max 2000 characters'),
        ('status',               'String',    'enum: active | deleted'),
    ]),
    ('Dish Collection', [
        ('name',                 'String',    'Required'),
        ('description',          'String',    ''),
        ('price',                'Number',    'In PKR'),
        ('dietary_tags',         '[String]',  'vegan | halal | gluten-free | vegetarian'),
        ('image_url',            'String',    ''),
        ('restaurant_id',        'ObjectId',  'Ref: Restaurant'),
    ]),
    ('Tag Collection', [
        ('name',                 'String',    'Unique'),
        ('name_lower',           'String',    'Normalised for search'),
        ('type',                 'String',    'enum: cuisine | dietary | feature'),
        ('status',               'String',    'enum: pending | approved'),
        ('created_by',           'ObjectId',  'Ref: User'),
        ('usage_count',          'Number',    'Denormalised counter'),
    ]),
]

for schema_title, fields in schemas:
    heading(schema_title, level=3, size=12, space_before=14)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style     = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, txt in enumerate(['Field', 'Type', 'Notes']):
        set_cell_bg(hdr[i], '1a1a2e')
        p   = hdr[i].paragraphs[0]
        r   = p.add_run(txt)
        r.font.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = WHITE
    for field, ftype, note in fields:
        row = tbl.add_row().cells
        pf  = row[0].paragraphs[0]
        rf  = pf.add_run(field)
        rf.font.bold = True; rf.font.size = Pt(10); rf.font.color.rgb = DARK
        pt  = row[1].paragraphs[0]
        rt  = pt.add_run(ftype)
        rt.font.size = Pt(9.5); rt.font.color.rgb = RGBColor(0x5b, 0x21, 0xb6)
        pn  = row[2].paragraphs[0]
        rn  = pn.add_run(note)
        rn.font.size = Pt(9.5); rn.font.color.rgb = GREY
        if note in ('Required', 'Unique index', 'Unique', '2dsphere index'):
            rn.font.color.rgb = RGBColor(0xe2, 0xb0, 0x4a)
            rn.font.bold      = True
        for cell in row:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after  = Pt(2)
    doc.add_paragraph()

heading('Key Database Indexes', level=3)
bullet([
    'Restaurant: 2dsphere index on location field for geospatial queries',
    'Restaurant: Compound text index on name + description for full-text search',
    'Restaurant: Unique compound index on name_normalized + address_normalized to prevent duplicates',
    'Review: Unique compound index on user_id + restaurant_id (one review per user per restaurant)',
    'ReviewVote: Unique compound index on review_id + user_id (one vote per user per review)',
    'RefreshToken: TTL index on expiresAt field for automatic cleanup of expired tokens',
])

divider()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. UI SCREENSHOTS
# ══════════════════════════════════════════════════════════════════════════════
section_label('Section 09')
heading('9. UI Screenshots', level=1)
body(
    'The application is built as a Single Page Application (SPA) using React 18 with React Router v6 '
    'for client-side navigation. All pages are fully responsive and styled with custom CSS.'
)

screenshots = [
    ('screenshots/home.jpg',        'Fig 1 — Home Page: Hero search section with Browse by Category tags'),
    ('screenshots/restaurants.jpg', 'Fig 2 — Restaurant Listing: Filter panel with cuisine, price, rating and tag filters'),
    ('screenshots/login.jpg',       'Fig 3 — Authentication: Login form with forgot password and sign-up links'),
]

for path, caption in screenshots:
    try:
        doc.add_picture(path, width=Inches(5.5))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap  = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr   = cap.runs[0]
        cr.font.size      = Pt(9)
        cr.font.italic    = True
        cr.font.color.rgb = GREY
        cap.paragraph_format.space_after = Pt(14)
    except Exception as e:
        body(f'[Screenshot not available: {path} — {e}]', italic=True)

divider()
p = doc.add_paragraph('FoodScope  ·  MERN Stack Web Application  ·  Advanced Web Technologies Lab Terminal')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.runs[0]
r.font.size = Pt(9); r.font.color.rgb = GREY; r.font.italic = True

doc.save('FoodScope_Project_Document.docx')
print('Document saved: FoodScope_Project_Document.docx')
